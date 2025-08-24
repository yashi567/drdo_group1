# 02_extract_clean_chunk.py
# Usage:
#   python 02_extract_clean_chunk.py --manifest data\\manifest.json --out data\\chunks.jsonl --chunk-size 1200 --overlap 200 [--pdf-progress-every 25]
#
# Produces data/chunks.jsonl where each line is:
#   {"id": "...", "text": "...", "metadata": {...}}
#
# Notes:
# - PDF: uses PyMuPDF (fitz) for reliable text extraction.
# - DOCX: python-docx
# - TXT/MD: plain read
# - HTML: BeautifulSoup to get visible text
# - Cleans whitespace, removes very common header/footer lines, creates overlapping char-based chunks.

import argparse
import json
import re
import uuid
import os
import time
from pathlib import Path
from collections import Counter

import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup

# Global, set from CLI for PDF progress cadence
PDF_PROGRESS_EVERY = 25

def log(msg: str):
    print(msg, flush=True)

def read_manifest(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)["files"]

def extract_text_pdf(fp: Path) -> str:
    text_parts = []
    try:
        with fitz.open(str(fp)) as doc:
            n_pages = len(doc)
            log(f"  → PDF opened: {n_pages} pages")
            for i, page in enumerate(doc, start=1):
                t = page.get_text("text")
                text_parts.append(t or "")
                if (i % PDF_PROGRESS_EVERY == 0) or (i == n_pages):
                    log(f"    ...processed page {i}/{n_pages}")
    except Exception as e:
        log(f"[WARN] PDF failed ({fp}): {e}")
    return "\n".join(text_parts).strip()

def extract_text_docx(fp: Path) -> str:
    try:
        doc = Document(str(fp))
        log(f"  → DOCX opened: {len(doc.paragraphs)} paragraphs")
        return "\n".join(p.text for p in doc.paragraphs).strip()
    except Exception as e:
        log(f"[WARN] DOCX failed ({fp}): {e}")
        return ""

def extract_text_txt(fp: Path) -> str:
    try:
        size = os.path.getsize(fp)
        log(f"  → TXT/MD read (~{size/1024:.1f} KB)")
        return fp.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as e:
        log(f"[WARN] TXT failed ({fp}): {e}")
        return ""

def extract_text_html(fp: Path) -> str:
    try:
        raw = fp.read_text(encoding="utf-8", errors="ignore")
        log(f"  → HTML read (~{len(raw)/1024:.1f} KB), parsing…")
        soup = BeautifulSoup(raw, "html.parser")
        for s in soup(["script", "style", "noscript"]):
            s.extract()
        text = soup.get_text(separator="\n")
        return text.strip()
    except Exception as e:
        log(f"[WARN] HTML failed ({fp}): {e}")
        return ""

def normalize_whitespace(text: str) -> str:
    # Standardize newlines and collapse excessive spaces
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    # Remove blank lines at extremes
    lines = [ln.strip() for ln in text.split("\n")]
    # Drop lines that are just page numbers (e.g., "12")
    lines = [ln for ln in lines if not re.fullmatch(r"\d{1,4}", ln)]
    # Remove runs of empty lines > 1
    out_lines = []
    last_blank = False
    for ln in lines:
        if ln == "":
            if not last_blank:
                out_lines.append(ln)
            last_blank = True
        else:
            out_lines.append(ln)
            last_blank = False
    return "\n".join(out_lines).strip()

def remove_common_headers_footers(text: str, threshold=5) -> str:
    # Heuristic: remove lines that repeat too often (headers/footers)
    lines = text.split("\n")
    freq = Counter([ln for ln in lines if ln and len(ln) < 120])  # short lines likely headers
    common = {ln for ln, c in freq.items() if c >= threshold}
    if not common:
        return text
    return "\n".join([ln for ln in lines if ln not in common])

def chunk_text(text: str, chunk_size=1200, overlap=200):
    """
    Character-based sliding window with overlap.
    Safer loop that guarantees forward progress even if sentence-boundary
    detection fails or suggests a non-advancing cut.
    """
    if not text:
        return []

    if overlap >= chunk_size:
        # Avoid infinite loops; enforce a sane overlap
        overlap = max(0, chunk_size // 4)

    chunks = []
    start = 0
    n = len(text)
    tail_span = 120  # look back window to find a sentence end

    while start < n:
        end = min(start + chunk_size, n)
        snippet = text[start:end]

        # Try to end at a sentence boundary within the last `tail_span` chars
        tail = snippet[-tail_span:] if len(snippet) > tail_span else snippet
        # Use a zero-width lookahead so we don't consume the capital/quote
        m = re.search(r"[.!?]\s+(?=[A-Z(“\"'])", tail)

        if m:
            # Map tail position back to absolute text position
            proposed_end = start + len(snippet) - (len(tail) - m.end())
            # Only accept if it actually moves forward
            if start < proposed_end <= n:
                end = proposed_end
                snippet = text[start:end]

        # Fallback: guarantee forward progress by at least 1 char
        if end <= start:
            end = min(start + max(1, chunk_size), n)
            snippet = text[start:end]

        chunks.append((start, end, snippet.strip()))

        if end >= n:
            break

        # Next window start with overlap, but ensure progress
        new_start = end - overlap
        if new_start <= start:
            new_start = start + 1
        start = new_start

    return chunks


def extract_any(fp: Path) -> str:
    ext = fp.suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(fp)
    elif ext == ".docx":
        return extract_text_docx(fp)
    elif ext in {".txt", ".md"}:
        return extract_text_txt(fp)
    elif ext in {".html", ".htm"}:
        return extract_text_html(fp)
    else:
        log(f"  → Skipping unsupported file type: {ext}")
        return ""

def main():
    global PDF_PROGRESS_EVERY

    ap = argparse.ArgumentParser()
    ap.add_argument("--manifest", required=True, help="Path to manifest.json from step 1")
    ap.add_argument("--out", default="data/chunks.jsonl", help="Output JSONL of chunks")
    ap.add_argument("--chunk-size", type=int, default=1200)
    ap.add_argument("--overlap", type=int, default=200)
    ap.add_argument("--pdf-progress-every", type=int, default=25, help="Log progress every N PDF pages")
    args = ap.parse_args()

    PDF_PROGRESS_EVERY = max(1, args.pdf_progress_every)

    start_time = time.time()
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    log(f"[INIT] manifest={args.manifest}  out={args.out}  chunk_size={args.chunk_size}  overlap={args.overlap}  pdf_progress_every={PDF_PROGRESS_EVERY}")
    files = read_manifest(Path(args.manifest))
    log(f"[INIT] Found {len(files)} file(s) in manifest")

    total = 0
    with out_path.open("w", encoding="utf-8") as out_f:
        for i, entry in enumerate(files, start=1):
            fp = Path(entry["path"])
            ext = fp.suffix.lower()
            exists = fp.exists()
            size_kb = (os.path.getsize(fp) / 1024.0) if exists else 0
            log(f"\n[{i}/{len(files)}] Processing: {fp}  (ext={ext}, size≈{size_kb:.1f} KB, exists={exists})")
            if not exists:
                log(f"  ! File not found, skipping.")
                continue

            t0 = time.time()
            text = extract_any(fp)
            t_extract = time.time() - t0
            log(f"  ✓ Extracted {len(text):,} characters in {t_extract:.2f}s")
            if not text:
                log("  ! Empty extraction, skipping.")
                continue

            t1 = time.time()
            text_norm = normalize_whitespace(text)
            log(f"  ✓ Normalized -> {len(text_norm):,} characters ({len(text_norm) - len(text):+d} vs raw) in {time.time()-t1:.2f}s")

            t2 = time.time()
            text_clean = remove_common_headers_footers(text_norm, threshold=5)
            log(f"  ✓ Headers/footers cleaned -> {len(text_clean):,} characters ({len(text_clean) - len(text_norm):+d}) in {time.time()-t2:.2f}s")

            t3 = time.time()
            chunks = chunk_text(text_clean, chunk_size=args.chunk_size, overlap=args.overlap)
            log(f"  ✓ Chunked into {len(chunks):,} chunks in {time.time()-t3:.2f}s")

            for idx, (s, e, chunk) in enumerate(chunks):
                cid = str(uuid.uuid4())
                rec = {
                    "id": cid,
                    "text": chunk,
                    "metadata": {
                        "source_path": str(fp),
                        "ext": ext,
                        "chunk_index": idx,
                        "start_char": s,
                        "end_char": e
                    }
                }
                out_f.write(json.dumps(rec, ensure_ascii=False) + "\n")
            total += len(chunks)
            log(f"  → Wrote {len(chunks)} chunk(s) for this file (total so far: {total:,})")

    elapsed = time.time() - start_time
    log(f"\n[OK] Wrote {total:,} chunks to {out_path} in {elapsed:.2f}s")

if __name__ == "__main__":
    main()
